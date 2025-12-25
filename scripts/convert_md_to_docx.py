import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by double newlines to form paragraphs/blocks
    blocks = content.split('\n\n')
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
            
        # Check if header
        if block.startswith('#'):
            # It's a header. Could be multiple lines if bad formatting, but usually one.
            lines = block.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    # Fallback for weird header wrapping
                    p = doc.add_paragraph(line)
        
        # Check if list
        elif block.startswith('* ') or block.startswith('- ') or block.lstrip().startswith('1. '):
            lines = block.split('\n')
            for line in lines:
                clean_line = line.strip()
                if clean_line.startswith('* ') or clean_line.startswith('- '):
                    p = doc.add_paragraph(clean_line[2:], style='List Bullet')
                    apply_formatting(p)
                elif clean_line[0].isdigit() and '. ' in clean_line:
                    # Simple ordered list detection
                    parts = clean_line.split('. ', 1)
                    if len(parts) > 1:
                        p = doc.add_paragraph(parts[1], style='List Number')
                        apply_formatting(p)
                else:
                    # Continuation of a list item or just text? Treat as list continuation for now
                    p = doc.add_paragraph(clean_line, style='List Continue') # 'List Continue' might not exist, fallback to Normal
                    apply_formatting(p)

        # Code block (simple detection)
        elif block.startswith('```'):
            lines = block.split('\n')
            for line in lines:
                if '```' not in line:
                    p = doc.add_paragraph()
                    p.add_run(line).font.name = 'Courier New'

        # Normal Text Paragraph
        else:
            # Join lines in the block to ensure smooth wrapping
            clean_text = block.replace('\n', ' ')
            p = doc.add_paragraph(clean_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            apply_formatting(p)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

def apply_formatting(paragraph):
    # Bold handling: **text**
    # This is a simple implementation; deeply nested robust parsing would require more code
    # We will split by '**'
    
    # Access the text of the paragraph we just added (it's in runs[0] usually)
    # But wait, we added the text in the constructor. Python-docx creates one run.
    # We need to rebuild the runs to handle bold.
    
    original_text = paragraph.text
    # Clear existing runs
    paragraph.clear()
    
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', original_text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                paragraph.add_run(part)

if __name__ == "__main__":
    md_file = r"d:\SDA Project\MindMate\docs\MindMate_Project_Report.md"
    docx_file = r"d:\SDA Project\MindMate\docs\MindMate_Project_Report_v3.docx"
    
    if os.path.exists(md_file):
        parse_markdown_to_docx(md_file, docx_file)
    else:
        print(f"Error: Could not find {md_file}")
